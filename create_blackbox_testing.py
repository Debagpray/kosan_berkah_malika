#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script untuk membuat Blackbox Testing Table untuk projek Kos Berkah Malika
Format: No, Fitur, Skenario Pengujian, Input yang Diharapkan, Output yang Diharapkan, Status (Pass/Fail)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_blackbox_testing_doc():
    # Create document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    # Title
    title = doc.add_heading('BLACKBOX TESTING TABLE\nKOS BERKAH MALIKA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle/Description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('Dokumen pengujian blackbox untuk semua fitur pada Sistem Informasi Reservasi dan Pembayaran Kos Berkah Malika')
    run.italic = True
    
    doc.add_paragraph()
    
    # Define all features and test cases based on code analysis
    test_cases = [
        # ==================== MODUL PENYEWA / USER ====================
        {
            "kategori": "A. MODUL PENYEWA (USER)",
            "test_cases": [
                # A1. Registrasi Akun
                {"no": "A1.1", "fitur": "Registrasi Akun Penyewa", 
                 "skenario": "User mendaftar dengan data lengkap dan valid", 
                 "input": "Username unik, nama lengkap, no HP (angka), password (min 6 karakter), konfirmasi password sama", 
                 "output": "Akun berhasil dibuat, redirect ke halaman login dengan pesan sukses", 
                 "status": ""},
                {"no": "A1.2", "fitur": "Registrasi Akun Penyewa", 
                 "skenario": "User mendaftar dengan username yang sudah terdaftar", 
                 "input": "Username yang sudah ada di database", 
                 "output": "Muncul pesan error 'Username sudah terdaftar'", 
                 "status": ""},
                {"no": "A1.3", "fitur": "Registrasi Akun Penyewa", 
                 "skenario": "User mendaftar dengan password dan konfirmasi tidak cocok", 
                 "input": "Password != konfirmasi password", 
                 "output": "Muncul pesan error 'Password tidak cocok'", 
                 "status": ""},
                {"no": "A1.4", "fitur": "Registrasi Akun Penyewa", 
                 "skenario": "User mendaftar dengan nomor HP mengandung huruf", 
                 "input": "No HP berisi karakter non-numerik", 
                 "output": "Muncul pesan error 'Nomor HP hanya boleh berisi angka'", 
                 "status": ""},
                {"no": "A1.5", "fitur": "Registrasi Akun Penyewa", 
                 "skenario": "User mendaftar dengan password kurang dari 6 karakter", 
                 "input": "Password panjang < 6 karakter", 
                 "output": "Muncul pesan error 'Password minimal 6 karakter'", 
                 "status": ""},
                
                # A2. Login Penyewa
                {"no": "A2.1", "fitur": "Login Penyewa", 
                 "skenario": "User login dengan kredensial valid", 
                 "input": "Username dan password yang sesuai dengan database", 
                 "output": "Login berhasil, redirect ke halaman utama (index.php)", 
                 "status": ""},
                {"no": "A2.2", "fitur": "Login Penyewa", 
                 "skenario": "User login dengan password salah", 
                 "input": "Username benar, password salah", 
                 "output": "Muncul pesan error login gagal", 
                 "status": ""},
                {"no": "A2.3", "fitur": "Login Penyewa", 
                 "skenario": "User login dengan username tidak terdaftar", 
                 "input": "Username tidak ada di database", 
                 "output": "Muncul pesan error login gagal", 
                 "status": ""},
                {"no": "A2.4", "fitur": "Login Penyewa", 
                 "skenario": "Admin mencoba login di halaman login penyewa", 
                 "input": "Username admin dengan role 'admin'", 
                 "output": "Muncul pesan 'Gunakan halaman Login Admin' dan session dihancurkan", 
                 "status": ""},
                
                # A3. Lihat Status Kamar
                {"no": "A3.1", "fitur": "Lihat Status Ketersediaan Kamar", 
                 "skenario": "User melihat status kamar yang tersedia", 
                 "input": "Kamar dengan status_kamar='tersedia' dan tidak ada reservasi aktif", 
                 "output": "Badge hijau 'Tersedia' ditampilkan", 
                 "status": ""},
                {"no": "A3.2", "fitur": "Lihat Status Ketersediaan Kamar", 
                 "skenario": "User melihat status kamar yang terisi", 
                 "input": "Kamar dengan reservasi status='Dikonfirmasi' dan tanggal saat ini dalam periode huni", 
                 "output": "Badge merah 'Terisi' ditampilkan", 
                 "status": ""},
                {"no": "A3.3", "fitur": "Lihat Status Ketersediaan Kamar", 
                 "skenario": "User melihat status kamar dalam perbaikan", 
                 "input": "Kamar dengan status_kamar='perbaikan'", 
                 "output": "Badge kuning 'Menunggu' ditampilkan", 
                 "status": ""},
                
                # A4. Detail Kamar
                {"no": "A4.1", "fitur": "Lihat Detail Kamar", 
                 "skenario": "User melihat detail kamar yang tersedia", 
                 "input": "ID kamar valid dengan status tersedia", 
                 "output": "Informasi lengkap kamar ditampilkan (foto, harga, fasilitas, deskripsi)", 
                 "status": ""},
                {"no": "A4.2", "fitur": "Lihat Detail Kamar", 
                 "skenario": "User melihat detail kamar yang terisi", 
                 "input": "ID kamar valid dengan status terisi", 
                 "output": "Informasi kamar ditampilkan dengan label 'Tidak Tersedia' dan tombol pesan disabled", 
                 "status": ""},
                {"no": "A4.3", "fitur": "Lihat Detail Kamar", 
                 "skenario": "User mengakses detail kamar dengan ID tidak valid", 
                 "input": "ID kamar tidak ada di database", 
                 "output": "Menampilkan pesan 'Kamar tidak ditemukan'", 
                 "status": ""},
                
                # A5. Pesan Kamar
                {"no": "A5.1", "fitur": "Form Pemesanan Kamar", 
                 "skenario": "User memesan kamar yang tersedia dengan data lengkap", 
                 "input": "Nama, no HP, tanggal masuk/keluar, durasi sewa, jumlah tamu, catatan", 
                 "output": "Reservasi dibuat dengan status 'Menunggu Pembayaran', popup Midtrans muncul", 
                 "status": ""},
                {"no": "A5.2", "fitur": "Form Pemesanan Kamar", 
                 "skenario": "User mencoba pesan kamar yang sudah terisi", 
                 "input": "ID kamar dengan status terisi atau dalam perbaikan", 
                 "output": "Muncul halaman error 'Maaf, Kamar Sudah Terisi' dengan opsi pilih kamar lain", 
                 "status": ""},
                {"no": "A5.3", "fitur": "Form Pemesanan Kamar", 
                 "skenario": "User input tanggal keluar sebelum/sama dengan tanggal masuk", 
                 "input": "tanggal_keluar <= tanggal_masuk", 
                 "output": "Muncul pesan error 'Tanggal keluar harus setelah tanggal masuk'", 
                 "status": ""},
                {"no": "A5.4", "fitur": "Form Pemesanan Kamar", 
                 "skenario": "User belum login mencoba akses halaman pemesanan", 
                 "input": "Akses users/pesan_kamar.php tanpa session user_id", 
                 "output": "Redirect ke halaman login dengan pesan 'Silakan login terlebih dahulu'", 
                 "status": ""},
                
                # A6. Pembayaran via Midtrans
                {"no": "A6.1", "fitur": "Pembayaran Online (Midtrans)", 
                 "skenario": "User melakukan pembayaran berhasil melalui Midtrans", 
                 "input": "Transaksi payment status = 'capture' atau 'settlement' dari Midtrans API", 
                 "output": "Status pembayaran update jadi 'lunas', status reservasi jadi 'Menunggu' konfirmansi admin", 
                 "status": ""},
                {"no": "A6.2", "fitur": "Pembayaran Online (Midtrans)", 
                 "skenario": "User menutup popup pembayaran tanpa menyelesaikan", 
                 "input": "User menutup popup Midtrans (onClose event)", 
                 "output": "Alert muncul, redirect ke halaman pesanan_saya.php dengan status pending", 
                 "status": ""},
                {"no": "A6.3", "fitur": "Pembayaran Online (Midtrans)", 
                 "skenario": "Pembayaran expired atau dibatalkan di sisi Midtrans", 
                 "input": "Transaction status = 'expire', 'cancel', atau 'deny'", 
                 "output": "Status reservasi update jadi 'Dibatalkan'", 
                 "status": ""},
                {"no": "A6.4", "fitur": "Pembayaran Online (Midtrans)", 
                 "skenario": "Auto-verify pembayaran pending saat user buka halaman pesanan", 
                 "input": "Halaman pesanan_saya.php dibuka dengan reservasi status 'Menunggu Pembayaran'", 
                 "output": "Sistem cek ke Midtrans API dan update status sesuai response", 
                 "status": ""},
                
                # A7. Pesanan Saya
                {"no": "A7.1", "fitur": "Riwayat Pesanan", 
                 "skenario": "User melihat daftar pesanan mereka", 
                 "input": "User login dengan id_pengguna memiliki reservasi", 
                 "output": "Daftar semua reservasi user ditampilkan dengan status dan detail", 
                 "status": ""},
                {"no": "A7.2", "fitur": "Riwayat Pesanan", 
                 "skenario": "User baru belum memiliki pesanan", 
                 "input": "User login tanpa record di tabel reservasi", 
                 "output": "Menampilkan pesan 'Anda belum memiliki pesanan'", 
                 "status": ""},
                {"no": "A7.3", "fitur": "Riwayat Pesanan", 
                 "skenario": "User akses halaman pesanan tanpa login", 
                 "input": "Akses users/pesanan_saya.php tanpa session", 
                 "output": "Redirect ke halaman login", 
                 "status": ""},
                
                # A8. Cetak Kwitansi
                {"no": "A8.1", "fitur": "Cetak Kwitansi Pembayaran", 
                 "skenario": "User cetak kwitansi untuk reservasi yang dikonfirmasi", 
                 "input": "ID reservasi dengan status_reservasi='Dikonfirmasi'", 
                 "output": "Halaman kwitansi terbuka dengan detail lengkap siap cetak/PDF", 
                 "status": ""},
                {"no": "A8.2", "fitur": "Cetak Kwitansi Pembayaran", 
                 "skenario": "User coba cetak kwitansi reservasi belum dikonfirmasi", 
                 "input": "ID reservasi dengan status selain 'Dikonfirmasi'", 
                 "output": "Redirect ke pesanan_saya.php atau pesan error", 
                 "status": ""},
                
                # A9. Form Data Penyewa
                {"no": "A9.1", "fitur": "Form Data Penyewa", 
                 "skenario": "User lengkapi data penyewa setelah reservasi dikonfirmasi", 
                 "input": "Data lengkap: KTP, alamat, pekerjaan, dll untuk reservasi yang sudah dikonfirmasi", 
                 "output": "Data tersimpan di tabel data_penyewa, status_huni='aktif'", 
                 "status": ""},
                {"no": "A9.2", "fitur": "Form Data Penyewa", 
                 "skenario": "User akses form tanpa reservasi dikonfirmasi", 
                 "input": "User tanpa reservasi status 'Dikonfirmasi' akses form", 
                 "output": "Pesan informasi bahwa perlu reservasi aktif terlebih dahulu", 
                 "status": ""},
                
                # A10. Notifikasi User
                {"no": "A10.1", "fitur": "Notifikasi Pengguna", 
                 "skenario": "User melihat notifikasi dari admin", 
                 "input": "User login dengan record di tabel notifikasi", 
                 "output": "Daftar notifikasi ditampilkan dengan status baca/belum dibaca", 
                 "status": ""},
                {"no": "A10.2", "fitur": "Notifikasi Pengguna", 
                 "skenario": "User belum memiliki notifikasi", 
                 "input": "User login tanpa record di tabel notifikasi", 
                 "output": "Menampilkan pesan 'Belum ada notifikasi'", 
                 "status": ""},
                
                # A11. Profil User
                {"no": "A11.1", "fitur": "Profil Pengguna", 
                 "skenario": "User melihat dan edit profil mereka", 
                 "input": "User login mengakses halaman profil", 
                 "output": "Data profil ditampilkan dan dapat diperbarui", 
                 "status": ""},
            ]
        },
        
        # ==================== MODUL ADMIN ====================
        {
            "kategori": "B. MODUL ADMIN",
            "test_cases": [
                # B1. Login Admin
                {"no": "B1.1", "fitur": "Login Admin", 
                 "skenario": "Admin login dengan kredensial valid", 
                 "input": "Username dan password admin yang valid dengan role='admin'", 
                 "output": "Login berhasil, redirect ke dashboard admin", 
                 "status": ""},
                {"no": "B1.2", "fitur": "Login Admin", 
                 "skenario": "Penyewa mencoba login di halaman admin", 
                 "input": "Username dengan role='penyewa'", 
                 "output": "Redirect ke halaman login admin atau pesan error", 
                 "status": ""},
                {"no": "B1.3", "fitur": "Login Admin", 
                 "skenario": "User belum login akses halaman admin", 
                 "input": "Akses admin/index.php tanpa session role='admin'", 
                 "output": "Redirect ke admin/login.php", 
                 "status": ""},
                
                # B2. Dashboard Admin
                {"no": "B2.1", "fitur": "Dashboard Admin", 
                 "skenario": "Admin melihat statistik sistem", 
                 "input": "Admin login mengakses dashboard", 
                 "output": "Statistik ditampilkan: total kamar, penghuni aktif, pesanan pending, total pendapatan", 
                 "status": ""},
                {"no": "B2.2", "fitur": "Dashboard Admin", 
                 "skenario": "Dashboard menampilkan reservasi terbaru", 
                 "input": "Ada record di tabel reservasi", 
                 "output": "Tabel 5 reservasi terbaru ditampilkan dengan info lengkap", 
                 "status": ""},
                
                # B3. Kelola Kamar
                {"no": "B3.1", "fitur": "Tambah Kamar Baru", 
                 "skenario": "Admin menambah kamar baru dengan data lengkap", 
                 "input": "Semua field required: nama_kamar, lantai, harga, foto, deskripsi, dll", 
                 "output": "Kamar berhasil ditambahkan ke database, redirect ke rooms.php", 
                 "status": ""},
                {"no": "B3.2", "fitur": "Edit Kamar", 
                 "skenario": "Admin mengedit data kamar yang ada", 
                 "input": "Perubahan pada field kamar (harga, deskripsi, status, dll)", 
                 "output": "Data kamar berhasil diperbarui", 
                 "status": ""},
                {"no": "B3.3", "fitur": "Hapus Kamar", 
                 "skenario": "Admin menghapus kamar yang tidak lagi tersedia", 
                 "input": "ID kamar valid, konfirmasi hapus", 
                 "output": "Kamar dihapus dari database, cascade ke reservasi terkait", 
                 "status": ""},
                {"no": "B3.4", "fitur": "Kelola Kamar - Status", 
                 "skenario": "Admin mengubah status kamar ke 'perbaikan'", 
                 "input": "Update status_kamar='perbaikan'", 
                 "output": "Kamar ditandai perbaikan, tidak bisa dipesan user", 
                 "status": ""},
                
                # B4. Kelola Reservasi
                {"no": "B4.1", "fitur": "Lihat Daftar Reservasi", 
                 "skenario": "Admin melihat semua reservasi yang masuk", 
                 "input": "Admin akses halaman bookings.php", 
                 "output": "Tabel semua reservasi ditampilkan dengan filter status", 
                 "status": ""},
                {"no": "B4.2", "fitur": "Konfirmasi Reservasi", 
                 "skenario": "Admin menyetujui reservasi yang menunggu", 
                 "input": "Update status_reservasi='Dikonfirmasi'", 
                 "output": "Status update, kamar jadi 'terisi', notifikasi dikirim ke user, data_penyewa auto-created", 
                 "status": ""},
                {"no": "B4.3", "fitur": "Batalkan Reservasi", 
                 "skenario": "Admin membatalkan reservasi", 
                 "input": "Update status_reservasi='Dibatalkan'", 
                 "output": "Status update, kamar jadi 'tersedia' kembali, notifikasi pembatalan dikirim", 
                 "status": ""},
                {"no": "B4.4", "fitur": "Edit Detail Reservasi", 
                 "skenario": "Admin mengedit tanggal masuk/keluar reservasi", 
                 "input": "Perubahan tanggal_masuk dan/atau tanggal_keluar", 
                 "output": "Detail reservasi diperbarui", 
                 "status": ""},
                
                # B5. Tambah Booking Offline
                {"no": "B5.1", "fitur": "Booking Offline (Walk-in)", 
                 "skenario": "Admin menambah reservasi untuk pembayaran tunai/langsung", 
                 "input": "Data pemesan, kamar, tanggal, durasi, metode pembayaran='Tunai'", 
                 "output": "Reservasi dibuat dengan status 'Dikonfirmasi', pembayaran tercatat 'lunas'", 
                 "status": ""},
                {"no": "B5.2", "fitur": "Booking Offline - Validasi", 
                 "skenario": "Admin coba booking offline untuk kamar tidak tersedia", 
                 "input": "ID kamar dengan status != 'tersedia'", 
                 "output": "Muncul error 'Kamar tidak tersedia'", 
                 "status": ""},
                
                # B6. Data Penyewa
                {"no": "B6.1", "fitur": "Lihat Data Penyewa Aktif", 
                 "skenario": "Admin melihat daftar penyewa yang sedang aktif", 
                 "input": "Filter status_huni='aktif'", 
                 "output": "Tabel penyewa aktif ditampilkan dengan detail kamar dan periode", 
                 "status": ""},
                {"no": "B6.2", "fitur": "Checkout Penyewa", 
                 "skenario": "Admin memproses penyewa yang akan keluar", 
                 "input": "Action checkout pada penyewa aktif", 
                 "output": "status_huni='sudah keluar', kamar jadi 'tersedia', reservasi jadi 'Selesai'", 
                 "status": ""},
                {"no": "B6.3", "fitur": "Hapus Data Penyewa", 
                 "skenario": "Admin menghapus record penyewa", 
                 "input": "Action delete pada penyewa, konfirmasi", 
                 "output": "Data penyewa dihapus, reservasi terkait diupdate jadi 'Selesai'", 
                 "status": ""},
                
                # B7. Kelola Pembayaran
                {"no": "B7.1", "fitur": "Lihat Riwayat Pembayaran", 
                 "skenario": "Admin melihat semua transaksi pembayaran", 
                 "input": "Admin akses halaman pembayaran.php", 
                 "output": "Tabel pembayaran ditampilkan dengan status dari Midtrans", 
                 "status": ""},
                {"no": "B7.2", "fitur": "Verifikasi Pembayaran Manual", 
                 "skenario": "Admin memverifikasi pembayaran yang masuk manual", 
                 "input": "Update status_transaksi='lunas' untuk pembayaran tertentu", 
                 "output": "Status pembayaran dan reservasi terkait diperbarui", 
                 "status": ""},
                {"no": "B7.3", "fitur": "Hapus Record Pembayaran", 
                 "skenario": "Admin menghapus record pembayaran yang tidak valid", 
                 "input": "Action delete pada pembayaran, konfirmasi", 
                 "output": "Record pembayaran dihapus dari database", 
                 "status": ""},
                
                # B8. Notifikasi
                {"no": "B8.1", "fitur": "Kirim Notifikasi ke User", 
                 "skenario": "Admin mengirim notifikasi ke pengguna tertentu", 
                 "input": "Pilih user, judul, pesan, jenis notifikasi", 
                 "output": "Notifikasi tersimpan di tabel notifikasi untuk user tersebut", 
                 "status": ""},
                {"no": "B8.2", "fitur": "Broadcast Notifikasi", 
                 "skenario": "Admin mengirim notifikasi ke semua penyewa", 
                 "input": "Judul, pesan, jenis untuk broadcast", 
                 "output": "Notifikasi terkirim ke semua user dengan role='penyewa'", 
                 "status": ""},
                {"no": "B8.3", "fitur": "Hapus Notifikasi", 
                 "skenario": "Admin menghapus notifikasi dari riwayat", 
                 "input": "Action delete pada notifikasi tertentu", 
                 "output": "Notifikasi dihapus dari database", 
                 "status": ""},
                
                # B9. Laporan
                {"no": "B9.1", "fitur": "Generate Laporan Keuangan", 
                 "skenario": "Admin generate laporan pendapatan periode tertentu", 
                 "input": "Filter tanggal mulai dan selesai", 
                 "output": "Laporan ditampilkan dengan total pendapatan dan detail transaksi", 
                 "status": ""},
                {"no": "B9.2", "fitur": "Export Laporan", 
                 "skenario": "Admin export laporan ke file eksternal", 
                 "input": "Action export pada laporan yang sudah difilter", 
                 "output": "File laporan terdownload (CSV/PDF/Excel tergantung implementasi)", 
                 "status": ""},
                
                # B10. Manajemen Pengguna
                {"no": "B10.1", "fitur": "Kelola Akun Pengguna", 
                 "skenario": "Admin melihat daftar semua akun penyewa", 
                 "input": "Admin akses halaman users.php", 
                 "output": "Tabel semua pengguna dengan role='penyewa' ditampilkan", 
                 "status": ""},
                {"no": "B10.2", "fitur": "Hapus Akun Pengguna", 
                 "skenario": "Admin menghapus akun penyewa yang melanggar aturan", 
                 "input": "Action delete pada user, konfirmasi (bukan akun sendiri)", 
                 "output": "Akun user dihapus dari database", 
                 "status": ""},
                {"no": "B10.3", "fitur": "Prevent Delete Self", 
                 "skenario": "Admin mencoba menghapus akun mereka sendiri", 
                 "input": "Action delete pada user_id yang sama dengan session admin", 
                 "output": "Error 'Tidak dapat menghapus akun sendiri'", 
                 "status": ""},
            ]
        },
        
        # ==================== FITUR UMUM ====================
        {
            "kategori": "C. FITUR UMUM & SISTEM",
            "test_cases": [
                # C1. Logout
                {"no": "C1.1", "fitur": "Logout", 
                 "skenario": "User/Admin logout dari sistem", 
                 "input": "Klik tombol logout", 
                 "output": "Session dihancurkan, redirect ke halaman login masing-masing", 
                 "status": ""},
                
                # C2. Responsive Design
                {"no": "C2.1", "fitur": "Tampilan Mobile-Friendly", 
                 "skenario": "Akses aplikasi dari perangkat mobile", 
                 "input": "Browser dengan viewport lebar kecil (<768px)", 
                 "output": "UI menyesuaikan, navigasi dan konten tetap usable", 
                 "status": ""},
                
                # C3. WhatsApp Integration
                {"no": "C3.1", "fitur": "Tombol WhatsApp Float", 
                 "skenario": "User klik tombol WhatsApp untuk chat", 
                 "input": "Klik floating button WhatsApp", 
                 "output": "WhatsApp Web/App terbuka dengan pesan template sudah terisi", 
                 "status": ""},
                
                # C4. Google Maps Integration
                {"no": "C4.1", "fitur": "Peta Lokasi Kos", 
                 "skenario": "User melihat lokasi kos di peta", 
                 "input": "Scroll ke section lokasi di homepage", 
                 "output": "Google Maps embed ditampilkan dengan pin lokasi yang benar", 
                 "status": ""},
                
                # C5. Session Security
                {"no": "C5.1", "fitur": "Keamanan Session", 
                 "skenario": "User mencoba akses halaman tanpa hak akses", 
                 "input": "Penyewa akses URL admin/* atau sebaliknya", 
                 "output": "Redirect ke halaman login atau halaman yang sesuai role", 
                 "status": ""},
                
                # C6. Password Hashing
                {"no": "C6.1", "fitur": "Enkripsi Password", 
                 "skenario": "Sistem menyimpan password user baru", 
                 "input": "Registrasi dengan password plain text", 
                 "output": "Password disimpan sebagai hash BCRYPT cost 12 di database", 
                 "status": ""},
                
                # C7. Midtrans Notification Handler
                {"no": "C7.1", "fitur": "Midtrans Notification Handler", 
                 "skenario": "Server Midtrans mengirim webhook notification", 
                 "input": "POST request dari Midtrans dengan status transaksi", 
                 "output": "Status pembayaran dan reservasi diupdate otomatis", 
                 "status": ""},
            ]
        }
    ]
    
    # Add introduction paragraph
    intro = doc.add_heading('1. Pendahuluan', level=1)
    p_intro = doc.add_paragraph()
    p_intro.add_run('Dokumen ini berisi tabel pengujian Blackbox Testing untuk seluruh fitur pada sistem Kos Berkah Malika. ')
    p_intro.add_run('Pengujian dilakukan dengan memberikan input tertentu dan memverifikasi output yang dihasilkan sesuai dengan spesifikasi.')
    
    doc.add_paragraph()
    
    # Add scope
    scope = doc.add_heading('2. Ruang Lingkup Pengujian', level=1)
    p_scope = doc.add_paragraph()
    run_scope = p_scope.add_run('Pengujian mencakup tiga modul utama:\n')
    p_scope.add_run('• Modul Penyewa (User) - Registrasi, Login, Pencarian Kamar, Pemesanan, Pembayaran, dll.\n')
    p_scope.add_run('• Modul Admin - Dashboard, Kelola Kamar, Reservasi, Penyewa, Pembayaran, Notifikasi, Laporan\n')
    p_scope.add_run('• Fitur Umum - Logout, Keamanan Session, Integrasi Midtrans, WhatsApp, Google Maps')
    for run in p_scope.runs[1:]:
        run.italic = True
    
    doc.add_paragraph()
    
    # Create tables for each category
    for category_data in test_cases:
        kategori = category_data['kategori']
        cases = category_data['test_cases']
        
        # Category heading
        cat_heading = doc.add_heading(kategori, level=1)
        
        # Create table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        headers = ['No', 'Fitur', 'Skenario Pengujian', 'Input yang Diharapkan', 'Output yang Diharapkan', 'Status (Pass/Fail)']
        
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Set column widths
        for row in table.rows:
            row.cells[0].width = Cm(1.5)   # No
            row.cells[1].width = Cm(3.5)   # Fitur
            row.cells[2].width = Cm(5.0)   # Skenario
            row.cells[3].width = Cm(4.5)   # Input
            row.cells[4].width = Cm(5.0)   # Output
            row.cells[5].width = Cm(2.0)   # Status
        
        # Add data rows
        for case in cases:
            row_cells = table.add_row().cells
            row_cells[0].text = case['no']
            row_cells[1].text = case['fitur']
            row_cells[2].text = case['skenario']
            row_cells[3].text = case['input']
            row_cells[4].text = case['output']
            row_cells[5].text = case['status']
            
            # Set vertical alignment to center
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph()
    
    # Add conclusion
    conclusion = doc.add_heading('3. Kesimpulan', level=1)
    p_conclusion = doc.add_paragraph()
    p_conclusion.add_run('Dokumen blackbox testing ini dapat digunakan sebagai acuan untuk melakukan pengujian sistem secara menyeluruh. ')
    p_conclusion.add_run('Setiap skenario pengujian harus dieksekusi dan hasilnya dicatat pada kolom Status (Pass/Fail). ')
    p_conclusion.add_run('Jika ditemukan ketidaksesuaian antara output aktual dengan output yang diharapkan, maka perlu dilakukan perbaikan pada kode sumber.')
    
    doc.add_paragraph()
    
    # Add notes
    notes = doc.add_heading('4. Catatan', level=1)
    p_notes = doc.add_paragraph()
    p_notes.add_run('• Kolom Status diisi selama proses pengujian berlangsung\n')
    p_notes.add_run('• Test case dapat ditambahkan sesuai dengan perkembangan fitur sistem\n')
    p_notes.add_run('• Dokumentasikan bug atau issue yang ditemukan pada lembar terpisah\n')
    p_notes.add_run('• Pastikan environment testing menggunakan database terpisah dari production')
    
    # Save document
    file_path = '/workspace/Blackbox_Testing_Kos_Berkah_Malika.docx'
    doc.save(file_path)
    
    print(f"✅ Dokumen berhasil dibuat: {file_path}")
    print(f"📊 Total kategori: {len(test_cases)}")
    total_cases = sum(len(cat['test_cases']) for cat in test_cases)
    print(f"📝 Total test cases: {total_cases}")
    
    return file_path

if __name__ == "__main__":
    create_blackbox_testing_doc()
